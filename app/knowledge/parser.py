"""Unified Document Parser supporting PDF, DOCX, TXT, Markdown, HTML, JSON, CSV, Code, and Git repos."""

import os
from pathlib import Path
from typing import Dict, List, Optional, Type
from app.core.exceptions import JarvisError
from app.core.logger import JarvisLogger
from app.knowledge.interfaces import DocumentParser
from app.knowledge.models import Document, DocumentFormat

logger = JarvisLogger.get_logger("document_parser")


class DocumentParseError(JarvisError):
    """Raised when document parsing fails."""
    pass


class UnifiedDocumentParser(DocumentParser):
    """Unified document parser for multi-format file ingestion."""

    FORMAT_MAP: Dict[str, DocumentFormat] = {
        ".pdf": DocumentFormat.PDF,
        ".docx": DocumentFormat.DOCX,
        ".doc": DocumentFormat.DOCX,
        ".txt": DocumentFormat.TXT,
        ".md": DocumentFormat.MARKDOWN,
        ".markdown": DocumentFormat.MARKDOWN,
        ".html": DocumentFormat.HTML,
        ".htm": DocumentFormat.HTML,
        ".json": DocumentFormat.JSON,
        ".csv": DocumentFormat.CSV,
        ".py": DocumentFormat.CODE,
        ".js": DocumentFormat.CODE,
        ".ts": DocumentFormat.CODE,
        ".java": DocumentFormat.CODE,
        ".cpp": DocumentFormat.CODE,
        ".c": DocumentFormat.CODE,
        ".h": DocumentFormat.CODE,
        ".cs": DocumentFormat.CODE,
        ".go": DocumentFormat.CODE,
        ".rs": DocumentFormat.CODE,
    }

    def can_parse(self, file_path: str) -> bool:
        """Returns True if file extension is recognized or path is readable text/dir."""
        if os.path.isdir(file_path):
            return True
        ext = Path(file_path).suffix.lower()
        return ext in self.FORMAT_MAP or ext == ""

    def parse(self, file_path: str) -> Document:
        """Parses a target file or repository path into a Document model."""
        p = Path(file_path)
        if not p.exists():
            raise DocumentParseError(f"Target file path does not exist: '{file_path}'.")

        logger.info(f"Parsing document '{file_path}'...")

        if p.is_dir():
            return self._parse_directory_repo(p)

        ext = p.suffix.lower()
        fmt = self.FORMAT_MAP.get(ext, DocumentFormat.TXT)

        try:
            if fmt == DocumentFormat.PDF:
                raw_text = self._parse_pdf(p)
            elif fmt == DocumentFormat.DOCX:
                raw_text = self._parse_docx(p)
            elif fmt == DocumentFormat.HTML:
                raw_text = self._parse_html(p)
            else:
                raw_text = self._parse_plain_text(p)

            file_size = p.stat().st_size
            doc = Document(
                file_path=str(p.resolve()),
                format=fmt,
                title=p.name,
                raw_content=raw_text,
                file_size_bytes=file_size,
                metadata={"extension": ext}
            )
            logger.info(f"Document '{p.name}' parsed successfully ({doc.char_count} chars).")
            return doc
        except Exception as e:
            logger.error(f"Failed to parse document '{file_path}': {e}")
            raise DocumentParseError(f"Failed to parse '{file_path}': {e}") from e

    def _parse_plain_text(self, p: Path) -> str:
        """Reads plain text, code, JSON, CSV, or Markdown files."""
        try:
            return p.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            return p.read_text(encoding="latin-1", errors="replace")

    def _parse_pdf(self, p: Path) -> str:
        """Parses PDF file text content."""
        try:
            import pypdf
            reader = pypdf.PdfReader(str(p))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except Exception:
            # Fallback to plain text read if pypdf unavailable
            return self._parse_plain_text(p)

    def _parse_docx(self, p: Path) -> str:
        """Parses Microsoft Word DOCX file text content."""
        try:
            import docx
            doc = docx.Document(str(p))
            paras = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paras)
        except Exception:
            return self._parse_plain_text(p)

    def _parse_html(self, p: Path) -> str:
        """Parses HTML file text content."""
        content = self._parse_plain_text(p)
        try:
            import re
            text = re.sub(r"<script.*?>.*?</script>", "", content, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<style.*?>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<.*?>", " ", text)
            return re.sub(r"\s+", " ", text).strip()
        except Exception:
            return content

    def _parse_directory_repo(self, repo_dir: Path) -> Document:
        """Aggregates readable code and markdown files from a Git repository or directory."""
        content_parts: List[str] = []
        for root, _, files in os.walk(repo_dir):
            if ".git" in root or "__pycache__" in root or ".venv" in root:
                continue
            for file in files:
                fp = Path(root) / file
                if self.can_parse(str(fp)):
                    try:
                        text = self._parse_plain_text(fp)
                        content_parts.append(f"--- File: {fp.relative_to(repo_dir)} ---\n{text}")
                    except Exception:
                        pass

        full_content = "\n\n".join(content_parts)
        return Document(
            file_path=str(repo_dir.resolve()),
            format=DocumentFormat.GIT_REPO,
            title=repo_dir.name,
            raw_content=full_content,
            file_size_bytes=len(full_content.encode("utf-8")),
            metadata={"is_directory": True}
        )
